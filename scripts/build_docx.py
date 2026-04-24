"""Build a Word (.docx) policy report from the markdown chapter drafts.

視覺風格參考：時代力量 2025 年度報告
- 主色黃：#F5E625（標題底線、accent bar、封面色塊）
- 主文色：深灰黑 #333333（取代藍色 heading）
- 副色：淺灰 #888 for meta text
- 結構：封面 → 一頁式摘要 → 目錄 → 各章 → 附錄

Usage:
    python -m scripts.build_docx                           # Ch.0 + Ch.1 + Ch.3 + Ch.4
    python -m scripts.build_docx ch00 ch01 ch02 ch03 ch04 ch05

Output: reports/bridge_report.docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

from docx import Document  # noqa: E402
from docx.enum.table import WD_ALIGN_VERTICAL  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.shared import Cm, Pt, RGBColor  # noqa: E402


# --- NPP 2025 visual palette ---
NPP_YELLOW = "F5E625"       # 主視覺黃（底線、accent bar、封面色塊）
NPP_YELLOW_LIGHT = "FFF4A3"  # 淺黃（內文 highlight 背景）
DARK_TEXT = "222222"        # 主文黑（不用藍色）
SOFT_GRAY = "888888"        # 副文灰
MEDIUM_GRAY = "555555"      # 次要標題灰
CODE_BG = "F2F2F2"          # 程式碼背景淺灰

# 字型策略：CJK 用「微軟正黑體 / Microsoft JhengHei」
# - Windows Word 預設有此字型
# - Mac Word 會 fallback 到系統 CJK（Heiti/PingFang），外觀相近
# - 兩端視覺一致性高
CJK_FONT_HINT = "Microsoft JhengHei"
LATIN_FONT = "Microsoft JhengHei"   # 拉丁字母也走正黑體，確保視覺統一
MONO_FONT = "Menlo"


# -------- Font & shading helpers --------

def _apply_font(run, font_name=LATIN_FONT, cjk_font=CJK_FONT_HINT,
                size_pt=None, bold=False, color_hex=None):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), cjk_font)
    rFonts.set(qn("w:cs"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold:
        run.font.bold = True
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)


def _set_cjk_font(run, size_pt=None, bold=False, color_hex=None):
    _apply_font(run, font_name=LATIN_FONT, cjk_font=CJK_FONT_HINT,
                size_pt=size_pt, bold=bold, color_hex=color_hex or DARK_TEXT)


def _add_yellow_bottom_border(paragraph, size=12, color_hex=NPP_YELLOW):
    """Add NPP-style yellow bar underneath a heading paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    # replace any existing bottom border
    existing = pBdr.find(qn("w:bottom"))
    if existing is not None:
        pBdr.remove(existing)
    pBdr.append(bottom)


def _set_paragraph_shading(paragraph, fill_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    existing = pPr.find(qn("w:shd"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(shd)


def _shade_cell(cell, color_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


# -------- Inline markup --------

INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)


def _add_inline(paragraph, text: str):
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_cjk_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            _apply_font(run, font_name=MONO_FONT, cjk_font=CJK_FONT_HINT,
                        size_pt=10, color_hex="884422")
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                run = paragraph.add_run(m.group(1))
                _set_cjk_font(run)
                suffix = paragraph.add_run(f" ({m.group(2)})")
                _set_cjk_font(suffix, size_pt=9, color_hex=SOFT_GRAY)
        else:
            run = paragraph.add_run(part)
            _set_cjk_font(run)


# -------- Block parser --------

def _parse_markdown_into_doc(doc: Document, md: str):
    lines = md.splitlines()
    i = 0
    buf: list[str] = []

    def flush():
        nonlocal buf
        if buf:
            text = " ".join(s.strip() for s in buf if s.strip())
            if text:
                p = doc.add_paragraph()
                _add_inline(p, text)
            buf = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.startswith("```"):
            flush()
            i += 1
            code = []
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            _set_paragraph_shading(p, CODE_BG)
            for cl in code:
                run = p.add_run(cl + "\n")
                _apply_font(run, font_name=MONO_FONT, cjk_font=CJK_FONT_HINT,
                            size_pt=9, color_hex="333333")
            continue

        # Horizontal rule
        if re.match(r"^\s*---+\s*$", line):
            flush()
            # NPP-style yellow divider
            sep = doc.add_paragraph()
            _add_yellow_bottom_border(sep, size=4)
            i += 1
            continue

        # Headings
        h_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if h_match:
            flush()
            level = len(h_match.group(1))
            heading_text = h_match.group(2).strip()
            _add_npp_heading(doc, heading_text, level)
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            flush()
            qt = line.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.right_indent = Cm(0.4)
            _set_paragraph_shading(p, "FFFCEB")
            # Yellow left bar via left border
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is None:
                pBdr = OxmlElement("w:pBdr")
                pPr.append(pBdr)
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "24")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), NPP_YELLOW)
            pBdr.append(left)
            _add_inline(p, qt)
            i += 1
            continue

        # Bulleted list
        bullet_match = re.match(r"^[\s]*[-*]\s+(.+)$", line)
        if bullet_match:
            flush()
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, bullet_match.group(1).strip())
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^[\s]*\d+\.\s+(.+)$", line)
        if num_match:
            flush()
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, num_match.group(1).strip())
            i += 1
            continue

        # Table
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            flush()
            tlines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tlines.append(lines[i].strip())
                i += 1
            _add_table(doc, tlines)
            continue

        if not line.strip():
            flush()
            i += 1
            continue

        buf.append(line)
        i += 1

    flush()


def _add_npp_heading(doc: Document, text: str, level: int):
    """Heading in NPP style: dark text, yellow underline accent for H1 & H2."""
    sizes = {1: 22, 2: 16, 3: 13, 4: 11}
    size = sizes.get(level, 11)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
    heading.paragraph_format.space_after = Pt(8)
    # Outline level for TOC
    pPr = heading._p.get_or_add_pPr()
    outlineLvl = OxmlElement("w:outlineLvl")
    outlineLvl.set(qn("w:val"), str(level - 1))
    pPr.append(outlineLvl)

    run = heading.add_run(text)
    _set_cjk_font(run, size_pt=size, bold=True, color_hex=DARK_TEXT)

    # NPP yellow bar under H1 and H2
    if level <= 2:
        _add_yellow_bottom_border(heading, size=16 if level == 1 else 12)


def _parse_table_row(row: str) -> list[str]:
    inner = row.strip().strip("|")
    return [c.strip() for c in inner.split("|")]


def _add_table(doc: Document, table_lines: list[str]):
    if not table_lines:
        return
    rows = []
    for line in table_lines:
        if re.match(r"^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line):
            continue
        rows.append(_parse_table_row(line))
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Light Grid Accent 1"

    for i, cells in enumerate(rows):
        for j in range(n_cols):
            cell = tbl.rows[i].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(cells[j] if j < len(cells) else "")
            _set_cjk_font(run, size_pt=9, bold=(i == 0))
            if i == 0:
                _shade_cell(cell, NPP_YELLOW)
            else:
                _shade_cell(cell, "FAFAFA" if i % 2 == 0 else "FFFFFF")


# -------- Cover & TOC --------

def _add_cover(doc: Document, chapters: list[str]):
    # Yellow top block
    block = doc.add_paragraph()
    block.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_shading(block, NPP_YELLOW)
    run = block.add_run("\u00a0")  # nbsp for height
    _set_cjk_font(run, size_pt=36, color_hex=NPP_YELLOW)

    # Blank spacer
    doc.add_paragraph()

    # Small eyebrow
    eyebrow = doc.add_paragraph()
    r = eyebrow.add_run("政策分析報告  |  時代力量辦公室提報")
    _set_cjk_font(r, size_pt=10, color_hex=SOFT_GRAY)

    # Main title
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("淡江大橋議題分析")
    _set_cjk_font(r, size_pt=28, bold=True, color_hex=DARK_TEXT)
    _add_yellow_bottom_border(title, size=24)

    # Subtitle
    sub = doc.add_paragraph()
    sub.paragraph_format.space_before = Pt(4)
    r = sub.add_run("從輿情資料到制度改革的 AI 輔助政策分析")
    _set_cjk_font(r, size_pt=14, color_hex=MEDIUM_GRAY)

    doc.add_paragraph()

    # Meta box
    meta = doc.add_paragraph()
    _set_paragraph_shading(meta, "FFFCEB")
    r = meta.add_run("資料 cutoff：2026-04-22    |    分析日期：2026-04-24\n"
                     f"本稿包含章節：{', '.join(chapters)}")
    _set_cjk_font(r, size_pt=10, color_hex=MEDIUM_GRAY)

    # Page break
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_toc_page(doc: Document, chapters: list[str], chapter_titles: dict[str, str]):
    """Simple manually-generated TOC (Word can auto-TOC but requires manual refresh)."""
    heading = doc.add_paragraph()
    r = heading.add_run("目錄")
    _set_cjk_font(r, size_pt=20, bold=True, color_hex=DARK_TEXT)
    _add_yellow_bottom_border(heading, size=16)

    doc.add_paragraph()

    for key in chapters:
        title = chapter_titles.get(key, key)
        p = doc.add_paragraph()
        # Chapter number (yellow box feel via size/color)
        num = key.replace("ch", "")
        rn = p.add_run(f"  {num}  ")
        _set_cjk_font(rn, size_pt=11, bold=True, color_hex=DARK_TEXT)
        r = p.add_run(f"  {title}")
        _set_cjk_font(r, size_pt=11, color_hex=DARK_TEXT)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# -------- Main --------

CHAPTER_FILES = {
    "ch00": ("執行摘要（一頁式）", "reports/ch00_summary.md"),
    "ch01": ("議題地景", "reports/ch01_draft.md"),
    "ch02": ("輿情結構", "reports/ch02_draft.md"),
    "ch03": ("AI 設計思路", "reports/ch03_draft.md"),
    "ch04": ("時代力量的切入機會", "reports/ch04_draft.md"),
    "ch05": ("展望 2026-05-12 之後", "reports/ch05_draft.md"),
}

DEFAULT_CHAPTERS = ["ch00", "ch01", "ch02", "ch03", "ch04", "ch05"]


def _configure_default_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = LATIN_FONT
    style.font.size = Pt(11)
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style.element.append(rPr)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), LATIN_FONT)
    rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    rFonts.set(qn("w:eastAsia"), CJK_FONT_HINT)
    rFonts.set(qn("w:cs"), LATIN_FONT)
    rPr.append(rFonts)


def main() -> int:
    chapters = sys.argv[1:] or DEFAULT_CHAPTERS
    unknown = [c for c in chapters if c not in CHAPTER_FILES]
    if unknown:
        print(f"Unknown chapters: {unknown}")
        print(f"Available: {list(CHAPTER_FILES.keys())}")
        return 2

    out_path = REPO_ROOT / "reports" / "bridge_report.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _configure_default_style(doc)

    # Cover
    _add_cover(doc, chapters)

    # Table of contents
    title_map = {k: v[0] for k, v in CHAPTER_FILES.items()}
    _add_toc_page(doc, chapters, title_map)

    # Chapters
    for idx, key in enumerate(chapters):
        _, rel_path = CHAPTER_FILES[key]
        md_path = REPO_ROOT / rel_path
        if not md_path.exists():
            p = doc.add_paragraph()
            r = p.add_run(f"[{title_map[key]}] 尚未起稿（檔案不存在：{rel_path}）")
            _set_cjk_font(r, size_pt=10, color_hex="AA0000")
            continue

        if idx > 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        _parse_markdown_into_doc(doc, md_path.read_text(encoding="utf-8"))

    doc.save(str(out_path))
    print(f"Word doc written: {out_path}")
    print(f"Open with: open {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
